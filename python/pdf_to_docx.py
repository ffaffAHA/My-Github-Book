import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import os
# from docx.shared import Inches
from io import BytesIO
def pdf_to_images(pdf_path):
    # 打开PDF文件
    pdf_document = fitz.open(pdf_path)
    images = []
    
    # 遍历PDF的每一页
    for page_number in range(len(pdf_document)):
        # 获取当前页
        page = pdf_document[page_number]
        # 将PDF页面渲染为图片
        pix = page.get_pixmap()
        # 保存图片到内存
        img_data = pix.tobytes("png")
        images.append(img_data)
    
    # 关闭PDF文件
    pdf_document.close()
    return images


def images_to_word(docx_path, images, pdf_filename):
    # 创建一个新的Word文档
    document = Document()
    
    # 添加PDF文件名作为标题
    document.add_heading(pdf_filename, level=1)
    
    # 遍历所有图片并将它们添加到Word文档中
    for img_data in images:
        image_stream = BytesIO(img_data)  # 将字节对象转换为BytesIO对象
        document.add_picture(image_stream, width=Inches(6.0))  # 设置图片宽度，可以根据需要调整
        document.add_page_break()  # 每张图片后添加分页符
    
    # 保存Word文档
    document.save(docx_path)
def process_pdfs_in_directory(pdf_directory, output_directory):
    # 遍历PDF目录中的所有PDF文件
    for pdf_filename in os.listdir(pdf_directory):
        if pdf_filename.endswith('.pdf'):
            pdf_path = os.path.join(pdf_directory, pdf_filename)
            images = pdf_to_images(pdf_path)
            docx_filename = pdf_filename[:-4] + '.docx'  # 替换扩展名
            docx_path = os.path.join(output_directory, docx_filename)
            images_to_word(docx_path, images, pdf_filename)
            print(f"已处理PDF文件：{pdf_filename}")

def main():
    pdf_directory = 'E:\新建文件夹'  # PDF文件所在目录
    output_directory = 'E:\新建文件夹'  # 输出Word文档所在目录
    process_pdfs_in_directory(pdf_directory, output_directory)
    print("所有PDF文件已处理完毕。")



if __name__ == "__main__":
    main()
